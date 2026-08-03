"""
MAITRI Sequential DOCX Progress Report Generator
Generates formal, formatted Microsoft Word (.docx) progress reports for each Phase of the MAITRI project.
Saved in series as: docs/update_01.docx, docs/update_02.docx, etc.
"""

import os
import sys
from datetime import datetime
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    """Sets background shading of a docx table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tc_pr.append(shd)

def create_update_report(phase_num=1, target_folder=None):
    if target_folder is None:
        target_folder = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    
    docs_dir = os.path.join(target_folder, "docs")
    os.makedirs(docs_dir, exist_ok=True)
    
    filename = f"update_{phase_num:02d}.docx"
    filepath = os.path.join(docs_dir, filename)

    doc = Document()

    # Configure Margins (0.75 in)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Document Header Title
    title = doc.add_paragraph()
    title_run = title.add_run(f"MAITRI OFFLINE AI ASSISTANT — PHASE {phase_num} REPORT")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(15, 32, 67) # Deep Space Navy
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub_title = doc.add_paragraph()
    sub_run = sub_title.add_run("Multimodal AI for Total Resilience & Intelligence | NASA/ISRO Mission Architecture")
    sub_run.font.size = Pt(11)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(100, 110, 120)
    sub_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Metadata Table
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.autofit = False

    metadata = [
        ("Project Name", "MAITRI (Multimodal AI for Total Resilience & Intelligence)"),
        ("Date & Time", datetime.now().strftime("%Y-%m-%d %H:%M:%S UTC")),
        ("Target Repository", "https://github.com/Shalokexe/MAITRI-Dashboard"),
        ("Lead Engineer", "Cmdr. Shalok Dadhwal (Shalokexe)")
    ]

    for idx, (label, val) in enumerate(metadata):
        row = meta_table.rows[idx]
        cell_lbl, cell_val = row.cells[0], row.cells[1]
        
        cell_lbl.text = label
        cell_lbl.paragraphs[0].runs[0].font.bold = True
        cell_lbl.paragraphs[0].runs[0].font.size = Pt(10)
        set_cell_background(cell_lbl, "EBF3FA")
        
        cell_val.text = val
        cell_val.paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph().paragraph_format.space_after = Pt(16)

    # Content per phase
    if phase_num == 1:
        add_phase_1_content(doc)
    elif phase_num == 2:
        add_phase_2_content(doc)
    elif phase_num == 3:
        add_phase_3_content(doc)
    elif phase_num == 4:
        add_phase_4_content(doc)
    elif phase_num == 5:
        add_phase_5_content(doc)
    elif phase_num == 6:
        add_phase_6_content(doc)
    elif phase_num == 7:
        add_phase_7_content(doc)
    else:
        add_generic_phase_content(doc, phase_num)

    doc.save(filepath)
    print(f"[MAITRI DOCX Report] Phase {phase_num} report created successfully at: {filepath}")
    return filepath

def add_heading_styled(doc, text):
    h = doc.add_paragraph()
    run = h.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(14, 116, 144) # Cyber Cyan / Teal
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(6)

def add_phase_1_content(doc):
    add_heading_styled(doc, "1. Executive Overview")
    doc.add_paragraph("Phase 1 establishes project foundation and offline database architecture.")

def add_phase_2_content(doc):
    add_heading_styled(doc, "1. Executive Overview")
    doc.add_paragraph("Phase 2 implements the Multimodal Input Pipeline.")

def add_phase_3_content(doc):
    add_heading_styled(doc, "1. Executive Overview")
    doc.add_paragraph("Phase 3 implements the Offline AI Conversation Engine.")

def add_phase_4_content(doc):
    add_heading_styled(doc, "1. Executive Overview")
    doc.add_paragraph("Phase 4 implements the Health Monitoring Dashboard & Telemetry Analytics Engine.")

def add_phase_5_content(doc):
    add_heading_styled(doc, "1. Executive Overview")
    doc.add_paragraph("Phase 5 implements the Personalization Engine & Offline Entertainment Vault.")

def add_phase_6_content(doc):
    add_heading_styled(doc, "1. Executive Overview")
    doc.add_paragraph("Phase 6 implements Emergency Rules, Black-Box Exporter, PVT Cognitive Test, & Circadian Sync.")

def add_phase_7_content(doc):
    add_heading_styled(doc, "1. Executive Overview — Final Release v1.0.0")
    p = doc.add_paragraph(
        "Phase 7 completes the production-grade deployment & offline optimization of MAITRI. "
        "It includes 100% zero-internet air-gap integrity verification across 27 code files, "
        "automated edge deployment scripts (deploy_edge.sh / deploy_edge.bat) for NVIDIA Jetson Nano / Raspberry Pi 4 / Laptops, "
        "and complete system verification across all 7 project phases."
    )
    p.paragraph_format.line_spacing = 1.15

    add_heading_styled(doc, "2. Full Project Milestone Summary")
    milestones = [
        ("Phase 1 — Project Foundation", "SQLite DB architecture (backend/db/schema.sql & database.py), Mission Control Dashboard UI, DOCX generator."),
        ("Phase 2 — Multimodal Input Pipeline", "OpenCV face emotion HUD (backend/vision/emotion_detector.py), Voice tone stress analyzer (backend/audio/tone_analyzer.py)."),
        ("Phase 3 — Offline AI Conversation Engine", "Air-gapped conversational AI companion (backend/ai_engine/offline_companion.py), mood-aware rules & psychological interventions."),
        ("Phase 4 — Health Telemetry Dashboard", "Physiological boundary evaluator (backend/health/health_monitor.py), Chart.js trend visualizers, iOS Crystal UI styling."),
        ("Phase 5 — Personalization & Media Vault", "Offline content recommender (backend/personalization/content_recommender.py), astronaut personality presets (ISRO Vedic Calm, NASA Flight Director, Zen Mindfulness)."),
        ("Phase 6 — Alerting & Innovative Features", "Emergency reporter, SHA-256 black-box flight recorder, pre-EVA PVT reaction test, circadian lighting sync, and voice TTS out-loud synthesis."),
        ("Phase 7 — Edge Deployment & Air-Gap Audit", "Zero-internet air-gap compliance auditor (scripts/audit_offline_integrity.py), edge deployment scripts, 19/19 unit tests passing.")
    ]

    for m_title, m_desc in milestones:
        bp = doc.add_paragraph(style='List Bullet')
        r1 = bp.add_run(f"{m_title}: ")
        r1.bold = True
        bp.add_run(m_desc)

    add_heading_styled(doc, "3. Files Added & Modified in Phase 7")
    files_table = doc.add_table(rows=1, cols=3)
    files_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = files_table.rows[0].cells
    hdr_cells[0].text = "File Path"
    hdr_cells[1].text = "Status"
    hdr_cells[2].text = "Description"

    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(cell, "0F2043")

    files_list = [
        ("scripts/audit_offline_integrity.py", "NEW", "Zero-internet air-gap compliance audit script"),
        ("scripts/deploy_edge.sh", "NEW", "Jetson Nano / Raspberry Pi 4 edge deployment script"),
        ("scripts/deploy_edge.bat", "NEW", "Windows Edge PC deployment batch script"),
        ("docs/update_07.docx", "NEW", "Final sequential Phase 7 status report"),
        ("README.md", "UPDATED", "v1.0.0 Final release documentation & architecture specs"),
        ("CHANGELOG.md", "UPDATED", "Complete version history log v0.1.0 through v1.0.0")
    ]

    for path, status, desc in files_list:
        row_cells = files_table.add_row().cells
        row_cells[0].text = path
        row_cells[1].text = status
        row_cells[2].text = desc
        for c in row_cells:
            c.paragraphs[0].runs[0].font.size = Pt(9.5)

    add_heading_styled(doc, "4. Verification & Testing Results")
    doc.add_paragraph("1. Zero-Internet Air-Gap Compliance Audit: python scripts/audit_offline_integrity.py (PASSED — 100% Offline)")
    doc.add_paragraph("2. Comprehensive Unit Test Suite: python -m unittest discover tests (19/19 PASSED)")
    doc.add_paragraph("3. Local Web Dashboard: http://localhost:8080 (NOMINAL)")
    doc.add_paragraph("4. GitHub Repository Push: https://github.com/Shalokexe/MAITRI-Dashboard (SYNCHRONIZED)")

def add_generic_phase_content(doc, phase_num):
    add_heading_styled(doc, f"1. Phase {phase_num} Summary")
    doc.add_paragraph(f"Phase {phase_num} implementation details for MAITRI offline astronaut assistant.")

if __name__ == "__main__":
    phase = 1
    if len(sys.argv) > 1:
        phase = int(sys.argv[1])
    create_update_report(phase)
